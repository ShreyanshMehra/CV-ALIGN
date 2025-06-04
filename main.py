from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
import os
import uuid
import pdfplumber
from docx import Document
from typing import Dict, Optional
import logging
from datetime import datetime
import re

from rag_engine import EnhancedRAGEngine, ChunkConfig

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Enhanced CV Scoring System", version="2.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class CVStore:
    def __init__(self):
        self.cvs: Dict[str, str] = {}
        self.jd: Optional[str] = None
        self.metadata: Dict[str, dict] = {}
        self.rag_engine = EnhancedRAGEngine(
            chunk_config=ChunkConfig(chunk_size=256, overlap=50, strategy="recursive")
        )
    
    def add_cv(self, cv_id: str, text: str, filename: str):
        self.cvs[cv_id] = text
        self.metadata[cv_id] = {
            "filename": filename,
            "upload_time": datetime.now().isoformat(),
            "text_length": len(text),
            "word_count": len(text.split())
        }
    
    def set_jd(self, text: str, filename: str):
        self.jd = text
        self.metadata["jd"] = {
            "filename": filename,
            "upload_time": datetime.now().isoformat(),
            "text_length": len(text),
            "word_count": len(text.split())
        }
    
    def get_status(self):
        return {
            "cvs_uploaded": len(self.cvs),
            "cv_ids": list(self.cvs.keys()),
            "jd_uploaded": self.jd is not None,
            "metadata": self.metadata
        }

cv_store = CVStore()

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + f"\n--- Page {page_num + 1} ---\n"
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
    except Exception as e:
        logger.error(f"Error reading PDF file: {e}")
        raise HTTPException(status_code=400, detail=f"Error reading PDF file: {str(e)}")
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error reading DOCX file: {e}")
        raise HTTPException(status_code=400, detail=f"Error reading DOCX file: {str(e)}")

def save_and_extract(file: UploadFile) -> tuple[str, str]:
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    ext = file.filename.split(".")[-1].lower()
    if ext not in ["pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    filename = f"{uuid.uuid4()}.{ext}"
    path = os.path.join(UPLOAD_DIR, filename)
    try:
        with open(path, "wb") as f:
            content = file.file.read()
            if len(content) == 0:
                raise HTTPException(status_code=400, detail="Empty file uploaded")
            f.write(content)
        if ext == "pdf":
            extracted_text = extract_text_from_pdf(path)
        else:
            extracted_text = extract_text_from_docx(path)
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the file")
        return extracted_text, filename
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
    finally:
        if os.path.exists(path):
            try:
                os.remove(path)
            except:
                pass

def extract_section(text: str, section_name: str) -> str:
    pattern = rf"{section_name}[\s:\-]*\n?(.*?)(?=\n[A-Z][a-zA-Z ]{{2,20}}[\s:\-]*\n|$)"
    match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
    return match.group(1).strip() if match else ""

import re

import re

def extract_cgpa(text):
    """
    Extracts the most likely CGPA (on a 10-point scale) from the text.
    Looks for patterns like 'CGPA: 9.1/10', 'CGPA 9.1', 'GPA 8.5 out of 10', etc.
    Returns the highest value found between 0 and 10.
    """
    cgpa_candidates = []

    # Patterns like 'CGPA: 9.1/10', 'CGPA 9.1', 'CGPA-9.1', 'GPA 8.5 out of 10'
    regex = r"(?:CGPA|GPA)[^\d]{0,5}(\d{1,2}\.\d{1,2})"
    for match in re.finditer(regex, text, re.IGNORECASE):
        val = float(match.group(1))
        if 0 < val <= 10:
            cgpa_candidates.append(val)

    # Also look for lines like 'CGPA: 9.1/10'
    regex2 = r"(?:CGPA|GPA)[^\d]{0,5}(\d{1,2}\.\d{1,2})\s*(?:/|out of)?\s*10"
    for match in re.finditer(regex2, text, re.IGNORECASE):
        val = float(match.group(1))
        if 0 < val <= 10:
            cgpa_candidates.append(val)

    # If nothing found, look for a number with '/10' nearby
    regex3 = r"(\d{1,2}\.\d{1,2})\s*/\s*10"
    for match in re.finditer(regex3, text):
        val = float(match.group(1))
        if 0 < val <= 10:
            cgpa_candidates.append(val)

    if cgpa_candidates:
        return max(cgpa_candidates)
    return None


def score_cgpa(cgpa, max_cgpa=10.0):
    if cgpa and 0 < cgpa <= max_cgpa:
        return round((cgpa / max_cgpa) * 100, 2)
    else:
        return 0.0




@app.get("/")
def read_root():
    return {"message": "Enhanced CV Scoring System API", "version": "2.1.0"}

@app.get("/status")
def get_status():
    return cv_store.get_status()

@app.post("/upload-cv/")
async def upload_cv(cv_id: str, file: UploadFile = File(...)):
    try:
        text, filename = save_and_extract(file)
        cv_store.add_cv(cv_id, text, filename)
        return {
            "message": f"CV {cv_id} uploaded successfully",
            "cv_id": cv_id,
            "filename": filename,
            "text_preview": text[:300] + "..." if len(text) > 300 else text,
            "word_count": len(text.split()),
            "character_count": len(text)
        }
    except Exception as e:
        logger.error(f"Error uploading CV {cv_id}: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-jd/")
async def upload_jd(file: UploadFile = File(...)):
    try:
        text, filename = save_and_extract(file)
        cv_store.set_jd(text, filename)
        return {
            "message": "Job Description uploaded successfully",
            "filename": filename,
            "text_preview": text[:300] + "..." if len(text) > 300 else text,
            "word_count": len(text.split()),
            "character_count": len(text)
        }
    except Exception as e:
        logger.error(f"Error uploading JD: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/score-cvs/")
async def score_cvs(background_tasks: BackgroundTasks):
    if not cv_store.jd:
        raise HTTPException(status_code=400, detail="Please upload a Job Description first")
    if not cv_store.cvs:
        raise HTTPException(status_code=400, detail="Please upload at least one CV")
    jd_word_count = len(cv_store.jd.split())
    if jd_word_count < 10:
        raise HTTPException(
            status_code=400,
            detail="Job Description is too short for meaningful analysis. Please upload a more detailed JD."
        )
    try:
        jd_keywords_with_weights = EnhancedRAGEngine.extract_jd_keywords_with_weights(cv_store.jd, top_n=20)
    except Exception as e:
        logger.error(f"Error extracting JD keywords: {e}")
        raise HTTPException(status_code=500, detail="Failed to extract keywords from JD.")

    # Section weights (must sum to 1.0)
    WEIGHTS = {
        "cv": 0.4,
        "projects": 0.2,
        "experience": 0.2,
        "cgpa": 0.1,
        "techskills": 0.1
    }

    results = []
    processing_errors = []

    for cv_id, cv_text in cv_store.cvs.items():
        try:
            logger.info(f"Processing CV: {cv_id}")
            cv_word_count = len(cv_text.split())
            if cv_word_count < 10:
                error_msg = f"CV {cv_id} is too short for analysis"
                processing_errors.append(error_msg)
                results.append({
                    "cv_id": cv_id,
                    "error": error_msg,
                    "filename": cv_store.metadata[cv_id]["filename"]
                })
                continue

            # --- Extract sections ---
            projects_text = extract_section(cv_text, "Projects")
            experience_text = extract_section(cv_text, "Experience")
            techskills_text = extract_section(cv_text, "Technical Skills")
            cgpa_val = extract_cgpa(cv_text)

            # --- Section-wise scoring ---
            def section_score(section_text):
                if not section_text or len(section_text.strip()) < 10:
                    return {
                        "semantic": 0.0,
                        "keyword": 0.0,
                        "combined": 0.0
                    }
                score_data = cv_store.rag_engine.calculate_comprehensive_score(section_text, cv_store.jd)
                semantic_score = score_data.get("final_score", 0.0) * 100
                keyword_score = EnhancedRAGEngine.score_cv_by_semantic_keywords(
                    section_text, jd_keywords_with_weights, cv_store.rag_engine.model, threshold=0.6
                )
                combined = EnhancedRAGEngine.combine_scores(semantic_score, keyword_score, 0.8, 0.2)
                return {
                    "semantic": float(round(semantic_score, 2)),
                    "keyword": float(round(keyword_score, 2)),
                    "combined": float(round(combined, 2))
                }

            # Section scores
            cv_scores = section_score(cv_text)
            projects_scores = section_score(projects_text)
            experience_scores = section_score(experience_text)
            techskills_scores = section_score(techskills_text)
            cgpa_score = score_cgpa(cgpa_val)

            # Weighted aggregation
            final_score = (
                WEIGHTS["cv"] * cv_scores["combined"] +
                WEIGHTS["projects"] * projects_scores["combined"] +
                WEIGHTS["experience"] * experience_scores["combined"] +
                WEIGHTS["cgpa"] * cgpa_score +
                WEIGHTS["techskills"] * techskills_scores["combined"]
            )

            # Determine match quality
            if final_score >= 80:
                match_quality = "Excellent Match"
            elif final_score >= 65:
                match_quality = "Good Match"
            elif final_score >= 50:
                match_quality = "Fair Match"
            else:
                match_quality = "Poor Match"

            cv_result = {
                "cv_id": cv_id,
                "filename": cv_store.metadata[cv_id]["filename"],
                "final_score": float(round(final_score, 2)),
                "match_quality": match_quality,
                "sections": {
                    "complete_cv": cv_scores,
                    "projects": projects_scores,
                    "experience": experience_scores,
                    "technical_skills": techskills_scores,
                    "cgpa": {
                        "value": cgpa_val,
                        "score": float(round(cgpa_score, 2))
                    }
                }
            }

            results.append(cv_result)

        except Exception as e:
            error_msg = f"Unexpected error processing CV {cv_id}: {str(e)}"
            logger.error(error_msg)
            processing_errors.append(error_msg)
            results.append({
                "cv_id": cv_id,
                "error": str(e),
                "filename": cv_store.metadata[cv_id]["filename"]
            })

    successful_results = [r for r in results if "error" not in r]
    successful_results.sort(key=lambda x: x.get("final_score", 0), reverse=True)

    # Relative scoring section for final_score
    if successful_results:
        max_score = successful_results[0]["final_score"]
        if max_score > 0:
            for r in successful_results:
                r["relative_score"] = round((r["final_score"] / max_score) * 100, 2)
        else:
            for r in successful_results:
                r["relative_score"] = 0.0


    return {
        "ranked_candidates": successful_results,
        "summary": {
            "total_cvs_processed": len(cv_store.cvs),
            "successful_processing": len(successful_results),
            "failed_processing": len(processing_errors),
            "highest_score": max([r.get("final_score", 0) for r in successful_results], default=0),
            "average_score": round(
                sum([r.get("final_score", 0) for r in successful_results]) / len(successful_results)
                if successful_results else 0, 2
            ),
            "processing_errors": processing_errors if processing_errors else None
        }
    }

    from fastapi import APIRouter

@app.post("/final-output/")
async def final_output():
    if not cv_store.jd:
        raise HTTPException(status_code=400, detail="Please upload a Job Description first")
    if not cv_store.cvs:
        raise HTTPException(status_code=400, detail="Please upload at least one CV")
    jd_word_count = len(cv_store.jd.split())
    if jd_word_count < 10:
        raise HTTPException(
            status_code=400,
            detail="Job Description is too short for meaningful analysis. Please upload a more detailed JD."
        )
    try:
        jd_keywords_with_weights = EnhancedRAGEngine.extract_jd_keywords_with_weights(cv_store.jd, top_n=20)
    except Exception as e:
        logger.error(f"Error extracting JD keywords: {e}")
        raise HTTPException(status_code=500, detail="Failed to extract keywords from JD.")

    WEIGHTS = {
        "cv": 0.4,
        "projects": 0.2,
        "experience": 0.2,
        "cgpa": 0.1,
        "techskills": 0.1
    }

    results = []
    for cv_id, cv_text in cv_store.cvs.items():
        try:
            projects_text = extract_section(cv_text, "Projects")
            experience_text = extract_section(cv_text, "Experience")
            techskills_text = extract_section(cv_text, "Technical Skills")
            cgpa_val = extract_cgpa(cv_text)

            def section_score(section_text):
                if not section_text or len(section_text.strip()) < 10:
                    return {
                        "semantic": 0.0,
                        "keyword": 0.0,
                        "combined": 0.0
                    }
                score_data = cv_store.rag_engine.calculate_comprehensive_score(section_text, cv_store.jd)
                semantic_score = score_data.get("final_score", 0.0) * 100
                keyword_score = EnhancedRAGEngine.score_cv_by_semantic_keywords(
                    section_text, jd_keywords_with_weights, cv_store.rag_engine.model, threshold=0.6
                )
                combined = EnhancedRAGEngine.combine_scores(semantic_score, keyword_score, 0.8, 0.2)
                return {
                    "semantic": float(round(semantic_score, 2)),
                    "keyword": float(round(keyword_score, 2)),
                    "combined": float(round(combined, 2))
                }

            cv_scores = section_score(cv_text)
            projects_scores = section_score(projects_text)
            experience_scores = section_score(experience_text)
            techskills_scores = section_score(techskills_text)
            cgpa_score = score_cgpa(cgpa_val)

            final_score = (
                WEIGHTS["cv"] * cv_scores["combined"] +
                WEIGHTS["projects"] * projects_scores["combined"] +
                WEIGHTS["experience"] * experience_scores["combined"] +
                WEIGHTS["cgpa"] * cgpa_score +
                WEIGHTS["techskills"] * techskills_scores["combined"]
            )

            cv_result = {
                "cv_id": cv_id,
                "filename": cv_store.metadata[cv_id]["filename"],
                "final_score": float(round(final_score, 2)),
                "sections": {
                    "cgpa": {
                        "value": cgpa_val,
                        "score": float(round(cgpa_score, 2))
                    },
                    "experience": experience_scores,
                    "projects": projects_scores,
                    "technical_skills": techskills_scores
                }
            }
            results.append(cv_result)
        except Exception as e:
            logger.error(f"Error processing CV {cv_id}: {e}")
            continue

    # --- Calculate relative section scores ---
    import numpy as np
    section_keys = ["cgpa", "experience", "projects", "technical_skills"]
    section_max = {}
    for key in section_keys:
        if key == "cgpa":
            scores = [r["sections"][key]["score"] for r in results]
        else:
            scores = [r["sections"][key]["combined"] for r in results]
        section_max[key] = max(scores) if scores else 1

    for r in results:
        r["relative_sections"] = {}
        for key in section_keys:
            if key == "cgpa":
                score = r["sections"][key]["score"]
            else:
                score = r["sections"][key]["combined"]
            max_score = section_max[key]
            rel_score = (score / max_score) * 100 if max_score > 0 else 0
            r["relative_sections"][key] = round(rel_score, 2)

    # --- Calculate overall relative score ---
    max_final_score = max([r["final_score"] for r in results], default=1)
    for r in results:
        r["relative_score"] = round((r["final_score"] / max_final_score) * 100, 2) if max_final_score > 0 else 0

    # --- Generate Pros and review ---
    def generate_pros_review_relative(cv_result):
        review = []
        rel = cv_result.get("relative_sections", {})
        if rel.get("cgpa", 0) >= 75:
            review.append("Good CGPA score.")
        elif rel.get("cgpa", 0) < 40:
            review.append("Low  CGPA score.")
        if rel.get("experience", 0) >= 75:
            review.append("Good industry experiences.")
        elif rel.get("experience", 0) < 40:
            review.append("Candidate doesn't have ample industry experience.")
        if rel.get("projects", 0) >= 75:
            review.append("Good projects in CV shows proficiency in solving real-world problems")
        elif rel.get("projects", 0) < 40:
            review.append("Projects section is not impressive.")
        if rel.get("technical_skills", 0) >= 75:
            review.append("Strong relative technical skills.")
        elif rel.get("technical_skills", 0) < 40:
            review.append("Weak relative technical skills.")
        if cv_result.get("relative_score", 0) >= 75:
            review.append("Overall the CV of this candidate looks good.")
        elif cv_result.get("relative_score", 0) < 40:
            review.append("The CV isn't the best out of all the other candidates.")
        return {"review": review, "review": review}

    for r in results:
        r["pros_review"] = generate_pros_review_relative(r)

    # --- Sort by relative_score descending ---
    results.sort(key=lambda x: x.get("relative_score", 0), reverse=True)

    return {
        "ranked_candidates": results,
        "summary": {
            "total_cvs_processed": len(results),
            "highest_relative_score": max([r.get("relative_score", 0) for r in results], default=0),
            "average_relative_score": round(
                sum([r.get("relative_score", 0) for r in results]) / len(results)
                if results else 0, 2
            ),
        }
    }


@app.delete("/clear-data/")
async def clear_data():
    cv_store.cvs.clear()
    cv_store.jd = None
    cv_store.metadata.clear()
    return {"message": "All data cleared successfully"}

@app.get("/cv/{cv_id}")
async def get_cv_details(cv_id: str):
    if cv_id not in cv_store.cvs:
        raise HTTPException(status_code=404, detail="CV not found")
    return {
        "cv_id": cv_id,
        "metadata": cv_store.metadata[cv_id],
        "text_preview": cv_store.cvs[cv_id][:500] + "..." if len(cv_store.cvs[cv_id]) > 500 else cv_store.cvs[cv_id]
    }
