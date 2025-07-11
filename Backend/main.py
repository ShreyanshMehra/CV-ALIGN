from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
import os
import uuid
import pdfplumber
from docx import Document
from typing import Dict, Optional
import logging
from datetime import datetime
import re

from Backend.rag_engine import EnhancedRAGEngine, ChunkConfig

def extract_section(text: str, section_name: str) -> str:
    """Extract text from a specific section of the CV."""
    patterns = [
        rf"{section_name}:?\s*\n",
        rf"\b{section_name}\b:?\s*\n",
        rf"\[{section_name}\]:?\s*\n",
        rf"<{section_name}>:?\s*\n"
    ]
    for pattern in patterns:
        matches = list(re.finditer(pattern, text, re.IGNORECASE))
        if matches:
            start = matches[0].end()
            next_section = None
            for next_pattern in [r"\n\s*[A-Z][A-Za-z\s]*:?\s*\n", r"\n\s*\[[A-Z][A-Za-z\s]*\]:?\s*\n"]:
                next_matches = list(re.finditer(next_pattern, text[start:]))
                if next_matches:
                    next_section = next_matches[0].start() + start
                    break
            section_text = text[start:next_section] if next_section else text[start:]
            return section_text.strip()
    return ""

def score_cgpa(cgpa_text: str) -> float:
    """Extract and score CGPA from text."""
    if not cgpa_text:
        return 0.0
    cgpa_patterns = [
        r"CGPA\s*[:/]?\s*(\d+\.\d+)",
        r"GPA\s*[:/]?\s*(\d+\.\d+)",
        r"CGPA/Percentage.*?(\d+\.\d+)",
        r"(?:B\.Tech|Bachelor|Degree).*?(\d+\.\d+)",
    ]
    for pattern in cgpa_patterns:
        matches = re.finditer(pattern, cgpa_text, re.IGNORECASE | re.DOTALL)
        for match in matches:
            try:
                cgpa = float(match.group(1))
                if cgpa <= 10.0:
                    return (cgpa / 10.0) * 100
            except ValueError:
                continue
    return 0.0

def extract_cgpa(text: str) -> float:
    cgpa_sections = [
        "Education", "Academic", "CGPA", "GPA", "B.Tech", "Bachelor"
    ]
    for section in cgpa_sections:
        section_text = extract_section(text, section)
        if section_text:
            cgpa = score_cgpa(section_text)
            if cgpa > 0:
                return cgpa
    return score_cgpa(text)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="Enhanced CV Scoring System", version="2.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allows all origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

class CVStore:
    def __init__(self):
        self.cvs: Dict[str, str] = {}
        self.jd: Optional[str] = None
        self.metadata: Dict[str, dict] = {}
        self.rag_engine = EnhancedRAGEngine(
            chunk_config=ChunkConfig(chunk_size=256, overlap=50, strategy="recursive")
        )
    def add_cv(self, cv_id: str, text: str, filename: str):
        self.cvs[cv_id] = text
        self.metadata[cv_id] = {
            "filename": filename,
            "upload_time": datetime.now().isoformat(),
            "text_length": len(text),
            "word_count": len(text.split())
        }
    def set_jd(self, text: str, filename: str):
        self.jd = text
        self.metadata["jd"] = {
            "filename": filename,
            "upload_time": datetime.now().isoformat(),
            "text_length": len(text),
            "word_count": len(text.split())
        }
    def get_status(self):
        return {
            "cvs_uploaded": len(self.cvs),
            "cv_ids": list(self.cvs.keys()),
            "jd_uploaded": self.jd is not None,
            "metadata": self.metadata
        }

cv_store = CVStore()

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + f"\n--- Page {page_num + 1} ---\n"
                except Exception as e:
                    logger.warning(f"Could not extract text from page {page_num + 1}: {e}")
    except Exception as e:
        logger.error(f"Error reading PDF file: {e}")
        raise HTTPException(status_code=400, detail=f"Error reading PDF file: {str(e)}")
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error reading DOCX file: {e}")
        raise HTTPException(status_code=400, detail=f"Error reading DOCX file: {str(e)}")

def save_and_extract(file: UploadFile) -> tuple[str, str]:
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")
    logger.info(f"Processing file: {file.filename}")
    ext = file.filename.split(".")[-1].lower()
    if ext not in ["pdf", "docx"]:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
    filename = f"{uuid.uuid4()}.{ext}"
    path = os.path.join(UPLOAD_DIR, filename)
    logger.info(f"Saving file to: {path}")
    try:
        with open(path, "wb") as f:
            content = file.file.read()
            if len(content) == 0:
                raise HTTPException(status_code=400, detail="Empty file uploaded")
            f.write(content)
            logger.info(f"File saved successfully: {len(content)} bytes")
        if ext == "pdf":
            logger.info("Extracting text from PDF")
            extracted_text = extract_text_from_pdf(path)
        else:
            logger.info("Extracting text from DOCX")
            extracted_text = extract_text_from_docx(path)
        if not extracted_text or not extracted_text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the file. Please ensure the file contains readable text.")
        logger.info(f"Successfully extracted {len(extracted_text)} characters of text")
        return extracted_text, filename
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing file: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")
    finally:
        if os.path.exists(path):
            try:
                os.remove(path)
                logger.info(f"Temporary file removed: {path}")
            except Exception as e:
                logger.error(f"Failed to remove temporary file {path}: {e}")
                pass

@app.get("/")
def read_root():
    return {"message": "Enhanced CV Scoring System API", "version": "2.1.0"}

@app.get("/status")
def get_status():
    return cv_store.get_status()

@app.post("/upload-cv")
async def upload_cv(cv_id: str, file: UploadFile = File(...)):
    try:
        logger.info(f"Received CV upload request - ID: {cv_id}, Filename: {file.filename}")
        if not file:
            raise HTTPException(status_code=400, detail="No file provided")
        if not cv_id:
            raise HTTPException(status_code=400, detail="No CV ID provided")
        first_byte = await file.read(1)
        if not first_byte:
            raise HTTPException(status_code=400, detail="File is empty")
        await file.seek(0)
        text, filename = save_and_extract(file)
        logger.info(f"Successfully extracted text from CV - ID: {cv_id}, Length: {len(text)}")
        cv_store.add_cv(cv_id, text, filename)
        logger.info(f"CV stored successfully - ID: {cv_id}")
        return {
            "message": f"CV {cv_id} uploaded successfully",
            "cv_id": cv_id,
            "filename": filename,
            "text_preview": text[:300] + "..." if len(text) > 300 else text,
            "word_count": len(text.split()),
            "character_count": len(text)
        }
    except HTTPException as he:
        logger.error(f"HTTP error uploading CV {cv_id}: {str(he)}")
        raise
    except Exception as e:
        logger.error(f"Error uploading CV {cv_id}: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-jd")
async def upload_jd(file: UploadFile = File(...)):
    try:
        text, filename = save_and_extract(file)
        cv_store.set_jd(text, filename)
        return {
            "message": "Job Description uploaded successfully",
            "filename": filename,
            "text_preview": text[:300] + "..." if len(text) > 300 else text,
            "word_count": len(text.split()),
            "character_count": len(text)
        }
    except Exception as e:
        logger.error(f"Error uploading JD: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/score-cvs/")
async def score_cvs(background_tasks: BackgroundTasks):
    if not cv_store.jd:
        raise HTTPException(status_code=400, detail="Please upload a Job Description first")
    if not cv_store.cvs:
        raise HTTPException(status_code=400, detail="Please upload at least one CV")
    jd_word_count = len(cv_store.jd.split())
    if jd_word_count < 10:
        raise HTTPException(
            status_code=400,
            detail="Job Description is too short for meaningful analysis. Please upload a more detailed JD."
        )
    try:
        jd_keywords_with_weights = EnhancedRAGEngine.extract_jd_keywords_with_weights(cv_store.jd, top_n=20)
    except Exception as e:
        logger.error(f"Error extracting JD keywords: {e}")
        raise HTTPException(status_code=500, detail="Failed to extract keywords from JD.")

    WEIGHTS = {
        "cv": 0.4,
        "projects": 0.2,
        "experience": 0.2,
        "cgpa": 0.1,
        "techskills": 0.1
    }

    results = []
    processing_errors = []

    for cv_id, cv_text in cv_store.cvs.items():
        try:
            logger.info(f"Processing CV: {cv_id}")
            cv_word_count = len(cv_text.split())
            if cv_word_count < 10:
                error_msg = f"CV {cv_id} is too short for analysis"
                processing_errors.append(error_msg)
                results.append({
                    "cv_id": cv_id,
                    "error": error_msg,
                    "filename": cv_store.metadata[cv_id]["filename"]
                })
                continue

            projects_text = extract_section(cv_text, "Projects")
            experience_text = extract_section(cv_text, "Experience")
            techskills_text = extract_section(cv_text, "Technical Skills")
            cgpa_val = extract_cgpa(cv_text)

            def section_score(section_text):
                if not section_text or len(section_text.strip()) < 10:
                    return {
                        "semantic": 0.0,
                        "keyword": 0.0,
                        "combined": 0.0
                    }
                score_data = cv_store.rag_engine.calculate_comprehensive_score(section_text, cv_store.jd)
                semantic_score = score_data.get("final_score", 0.0) * 100
                keyword_score = EnhancedRAGEngine.score_cv_by_semantic_keywords(
                    section_text, jd_keywords_with_weights, cv_store.rag_engine.get_embedding, threshold=0.6
                )
                combined = EnhancedRAGEngine.combine_scores(semantic_score, keyword_score, 0.8, 0.2)
                return {
                    "semantic": float(round(semantic_score, 2)),
                    "keyword": float(round(keyword_score, 2)),
                    "combined": float(round(combined, 2))
                }

            cv_scores = section_score(cv_text)
            projects_scores = section_score(projects_text)
            experience_scores = section_score(experience_text)
            techskills_scores = section_score(techskills_text)

            final_score = (
                WEIGHTS["cv"] * cv_scores["combined"] +
                WEIGHTS["projects"] * projects_scores["combined"] +
                WEIGHTS["experience"] * experience_scores["combined"] +
                WEIGHTS["cgpa"] * cgpa_val +
                WEIGHTS["techskills"] * techskills_scores["combined"]
            )

            if final_score >= 80:
                match_quality = "Excellent Match"
            elif final_score >= 65:
                match_quality = "Good Match"
            elif final_score >= 50:
                match_quality = "Fair Match"
            else:
                match_quality = "Poor Match"

            cv_result = {
                "cv_id": cv_id,
                "filename": cv_store.metadata[cv_id]["filename"],
                "final_score": float(round(final_score, 2)),
                "match_quality": match_quality,
                "sections": {
                    "complete_cv": cv_scores,
                    "projects": projects_scores,
                    "experience": experience_scores,
                    "technical_skills": techskills_scores,
                    "cgpa": {
                        "value": cgpa_val,
                        "score": float(round(cgpa_val, 2))
                    }
                }
            }

            results.append(cv_result)

        except Exception as e:
            error_msg = f"Unexpected error processing CV {cv_id}: {str(e)}"
            logger.error(error_msg)
            processing_errors.append(error_msg)
            results.append({
                "cv_id": cv_id,
                "error": str(e),
                "filename": cv_store.metadata[cv_id]["filename"]
            })

    successful_results = [r for r in results if "error" not in r]
    successful_results.sort(key=lambda x: x.get("final_score", 0), reverse=True)

    if successful_results:
        max_score = successful_results[0]["final_score"]
        if max_score > 0:
            for r in successful_results:
                r["relative_score"] = round((r["final_score"] / max_score) * 100, 2)
        else:
            for r in successful_results:
                r["relative_score"] = 0.0

    return {
        "ranked_candidates": successful_results,
        "summary": {
            "total_cvs_processed": len(cv_store.cvs),
            "successful_processing": len(successful_results),
            "failed_processing": len(processing_errors),
            "highest_score": max([r.get("final_score", 0) for r in successful_results], default=0),
            "average_score": round(
                sum([r.get("final_score", 0) for r in successful_results]) / len(successful_results)
                if successful_results else 0, 2
            ),
            "processing_errors": processing_errors if processing_errors else None
        }
    }

@app.post("/final-output/")
async def final_output():
    if not cv_store.jd:
        raise HTTPException(status_code=400, detail="Please upload a Job Description first")
    if not cv_store.cvs:
        raise HTTPException(status_code=400, detail="Please upload at least one CV")
    jd_word_count = len(cv_store.jd.split())
    if jd_word_count < 10:
        raise HTTPException(
            status_code=400,
            detail="Job Description is too short for meaningful analysis. Please upload a more detailed JD."
        )
    try:
        jd_keywords_with_weights = EnhancedRAGEngine.extract_jd_keywords_with_weights(cv_store.jd, top_n=20)
    except Exception as e:
        logger.error(f"Error extracting JD keywords: {e}")
        raise HTTPException(status_code=500, detail="Failed to extract keywords from JD.")

    WEIGHTS = {
        "cv": 0.4,
        "projects": 0.2,
        "experience": 0.2,
        "cgpa": 0.1,
        "techskills": 0.1
    }

    results = []
    for cv_id, cv_text in cv_store.cvs.items():
        try:
            projects_text = extract_section(cv_text, "Projects")
            experience_text = extract_section(cv_text, "Experience")
            techskills_text = extract_section(cv_text, "Technical Skills")
            cgpa_val = extract_cgpa(cv_text)

            def section_score(section_text):
                if not section_text or len(section_text.strip()) < 10:
                    return {
                        "semantic": 0.0,
                        "keyword": 0.0,
                        "combined": 0.0
                    }
                score_data = cv_store.rag_engine.calculate_comprehensive_score(section_text, cv_store.jd)
                semantic_score = score_data.get("final_score", 0.0) * 100
                keyword_score = EnhancedRAGEngine.score_cv_by_semantic_keywords(
                    section_text, jd_keywords_with_weights, cv_store.rag_engine.get_embedding, threshold=0.6
                )
                combined = EnhancedRAGEngine.combine_scores(semantic_score, keyword_score, 0.8, 0.2)
                return {
                    "semantic": float(round(semantic_score, 2)),
                    "keyword": float(round(keyword_score, 2)),
                    "combined": float(round(combined, 2))
                }

            cv_scores = section_score(cv_text)
            projects_scores = section_score(projects_text)
            experience_scores = section_score(experience_text)
            techskills_scores = section_score(techskills_text)

            final_score = (
                WEIGHTS["cv"] * cv_scores["combined"] +
                WEIGHTS["projects"] * projects_scores["combined"] +
                WEIGHTS["experience"] * experience_scores["combined"] +
                WEIGHTS["cgpa"] * cgpa_val +
                WEIGHTS["techskills"] * techskills_scores["combined"]
            )

            cv_result = {
                "cv_id": cv_id,
                "filename": cv_store.metadata[cv_id]["filename"],
                "final_score": float(round(final_score, 2)),
                "sections": {
                    "cgpa": {
                        "value": cgpa_val,
                        "score": float(round(cgpa_val, 2))
                    },
                    "experience": experience_scores,
                    "projects": projects_scores,
                    "technical_skills": techskills_scores
                }
            }
            results.append(cv_result)
        except Exception as e:
            logger.error(f"Error processing CV {cv_id}: {e}")
            continue

    import numpy as np
    section_keys = ["cgpa", "experience", "projects", "technical_skills"]
    section_max = {}
    for key in section_keys:
        if key == "cgpa":
            scores = [r["sections"][key]["score"] for r in results]
        else:
            scores = [r["sections"][key]["combined"] for r in results]
        section_max[key] = max(scores) if scores else 1

    for r in results:
        r["relative_sections"] = {}
        for key in section_keys:
            if key == "cgpa":
                score = r["sections"][key]["score"]
            else:
                score = r["sections"][key]["combined"]
            max_score = section_max[key]
            rel_score = (score / max_score) * 100 if max_score > 0 else 0
            r["relative_sections"][key] = round(rel_score, 2)

    max_final_score = max([r["final_score"] for r in results], default=1)
    for r in results:
        r["relative_score"] = round((r["final_score"] / max_final_score) * 100, 2) if max_final_score > 0 else 0

    def generate_pros_review_relative(cv_result):
        review = []
        rel = cv_result.get("relative_sections", {})
        if rel.get("cgpa", 0) >= 75:
            review.append("Good CGPA score.")
        elif rel.get("cgpa", 0) < 40:
            review.append("Low  CGPA score.")
        if rel.get("experience", 0) >= 75:
            review.append("Good industry experiences.")
        elif rel.get("experience", 0) < 40:
            review.append("Candidate doesn't have ample industry experience.")
        if rel.get("projects", 0) >= 75:
            review.append("Good projects in CV shows proficiency in solving real-world problems")
        elif rel.get("projects", 0) < 40:
            review.append("Projects section is not impressive.")
        if rel.get("technical_skills", 0) >= 75:
            review.append("Strong relative technical skills.")
        elif rel.get("technical_skills", 0) < 40:
            review.append("Weak relative technical skills.")
        if cv_result.get("relative_score", 0) >= 75:
            review.append("Overall the CV of this candidate looks good.")
        elif cv_result.get("relative_score", 0) < 40:
            review.append("The CV isn't the best out of all the other candidates.")
        return {"review": review, "review": review}

    for r in results:
        r["pros_review"] = generate_pros_review_relative(r)

    results.sort(key=lambda x: x.get("relative_score", 0), reverse=True)

    return {
        "ranked_candidates": results,
        "summary": {
            "total_cvs_processed": len(results),
            "highest_relative_score": max([r.get("relative_score", 0) for r in results], default=0),
            "average_relative_score": round(
                sum([r.get("relative_score", 0) for r in results]) / len(results)
                if results else 0, 2
            ),
        }
    }

@app.delete("/clear-data/")
async def clear_data():
    cv_store.cvs.clear()
    cv_store.jd = None
    cv_store.metadata.clear()
    return {"message": "All data cleared successfully"}

@app.get("/cv/{cv_id}")
async def get_cv_details(cv_id: str):
    if cv_id not in cv_store.cvs:
        raise HTTPException(status_code=404, detail="CV not found")
    return {
        "cv_id": cv_id,
        "metadata": cv_store.metadata[cv_id],
        "text_preview": cv_store.cvs[cv_id][:500] + "..." if len(cv_store.cvs[cv_id]) > 500 else cv_store.cvs[cv_id]
    }
